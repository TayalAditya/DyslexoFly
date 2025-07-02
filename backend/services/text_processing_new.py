from PyPDF2 import PdfReader
from docx import Document
import os
import magic  # Added for file type detection
import re

def extract_text_from_pdf(file_path):
    """Extract text from PDF files"""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except:
                continue
        
        if len(text.strip()) < 50:
            return "Error: This PDF appears to be image-based. Please convert to text-based PDF or use a different file format."
            
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return f"Error extracting PDF: {str(e)}"

def extract_text_from_docx(file_path):
    """Enhanced DOCX extraction with tables/headers"""
    try:
        doc = Document(file_path)
        text = []
        
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text.append(paragraph.text)
        
        # Headers/Footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    text.append(header.text)
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    text.append(footer.text)
                    
        return "\n".join(text)
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

def extract_text(file_path):
    """Enhanced text extraction with PNG/JPEG error handling"""
    try:
        # Verify actual file type
        file_type = magic.from_file(file_path, mime=True)
        _, extension = os.path.splitext(file_path)
        extension = extension.lower()
        
        # Handle based on actual type
        if 'pdf' in file_type or extension == '.pdf':
            return extract_text_from_pdf(file_path)
            
        elif 'document' in file_type or extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
            
        elif 'text' in file_type or extension == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
                
        elif 'image' in file_type or extension in ['.png', '.jpg', '.jpeg']:
            return "Error: Image files (PNG/JPEG) require OCR processing which is not currently supported. Please convert your image to text using an online OCR tool, or upload a text-based PDF, DOCX, or TXT file instead."
            
        else:
            return "Error: Unsupported file format. Please upload PDF, DOCX, or TXT files."
            
    except Exception as e:
        print(f"Text extraction error: {e}")
        return f"Error extracting text: {str(e)}"

def estimate_processing_time(text_content):
    """Estimate processing time based on content length"""
    if not text_content:
        return 10
    
    word_count = len(text_content.split())
    
    # Base time estimates (in seconds)
    base_time = 5
    word_processing_time = word_count * 0.02  # 0.02 seconds per word
    ai_processing_time = min(30, word_count * 0.05)  # AI processing, capped at 30s
    
    total_time = base_time + word_processing_time + ai_processing_time
    return int(total_time)

def get_text_statistics(text_content):
    """Get comprehensive text statistics"""
    if not text_content:
        return {"words": 0, "characters": 0, "sentences": 0, "paragraphs": 0}
    
    # Count words
    words = len(text_content.split())
    
    # Count characters
    characters = len(text_content)
    
    # Count sentences (basic sentence detection)
    sentences = len(re.findall(r'[.!?]+', text_content))
    
    # Count paragraphs
    paragraphs = len([p for p in text_content.split('\n') if p.strip()])
    
    # Reading time estimate (average 200 words per minute)
    reading_time_minutes = max(1, words // 200)
    
    return {
        "words": words,
        "characters": characters,
        "sentences": sentences,
        "paragraphs": paragraphs,
        "reading_time_minutes": reading_time_minutes,
        "complexity": "Simple" if words < 500 else "Medium" if words < 2000 else "Complex"
    }
