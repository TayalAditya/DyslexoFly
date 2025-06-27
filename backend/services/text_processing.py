from PyPDF2 import PdfReader
from docx import Document
import pytesseract
from PIL import Image
import os
import magic  # Added for file type detection
import re

def is_scanned_pdf(file_path):
    """Check if PDF is image-based/scanned"""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            if '/Font' not in page['/Resources']:
                return True
        return False
    except:
        return False

def extract_text_from_pdf(file_path):
    """Extract text from PDF files with hybrid approach"""
    try:
        # First attempt: Standard text extraction
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except:
                continue
        
        # Fallback for scanned PDFs
        if len(text.strip()) < 100 and is_scanned_pdf(file_path):
            return extract_text_from_image(file_path, is_pdf=True)
            
        return text
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Enhanced DOCX extraction with tables/headers"""
    try:
        doc = Document(file_path)
        text = []
        
        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            text.append(paragraph.text)
        
        # Headers/Footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    text.append(header.text)
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    text.append(footer.text)
                    
        return "\n".join(text)
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""

def extract_text_from_image(file_path, is_pdf=False):
    """Enhanced OCR with preprocessing"""
    try:
        # Preprocessing
        image = Image.open(file_path)
        
        if is_pdf:
            # Convert PDF page to image
            import pdf2image
            images = pdf2image.convert_from_path(file_path)
            image = images[0] if images else image
        
        # Preprocessing pipeline
        image = image.convert('L')  # Grayscale
        image = image.point(lambda x: 0 if x < 140 else 255)  # Binarization
        
        # OCR with language detection
        try:
            from langdetect import detect
            lang = detect(open(file_path, 'rb').read(1000)) or 'eng'
        except:
            lang = 'eng'
            
        # Custom config for better accuracy
        config = r'--oem 3 --psm 6'
        text = pytesseract.image_to_string(image, lang=lang, config=config)
        
        return text
    except Exception as e:
        print(f"Image OCR error: {e}")
        return ""

def extract_text(file_path):
    """Robust text extraction with verification"""
    try:
        # Verify actual file type
        file_type = magic.from_file(file_path, mime=True)
        _, extension = os.path.splitext(file_path)
        extension = extension.lower()
        
        # Handle based on actual type
        if 'pdf' in file_type or extension == '.pdf':
            return extract_text_from_pdf(file_path)
            
        elif 'document' in file_type or extension in ['.docx', '.doc']:
            return extract_text_from_docx(file_path)
            
        elif 'text' in file_type or extension == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
                
        elif 'image' in file_type or extension in ['.png', '.jpg', '.jpeg']:
            return extract_text_from_image(file_path)
            
        else:
            return "Unsupported file format"
            
    except Exception as e:
        print(f"General extraction error: {e}")
        return ""